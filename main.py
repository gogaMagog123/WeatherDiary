import requests
from bs4 import BeautifulSoup
from pprint import pprint
from docx import Document

import argparse


PAGE_URL = "https://arhivpogodi.ru/arhiv/sankt-peterburg/"

WIND_DIRECTIONS = [
    "С",
    "СВ",
    "В",
    "ЮВ",
    "Ю",
    "ЮЗ",
    "З",
    "СЗ",
]

parser = argparse.ArgumentParser(description='month and year')
parser.add_argument('month', help='month')
parser.add_argument('year', help='year')
args = parser.parse_args()

MONTH = args.month
YEAR = args.year

def generate_word_doc(data, filename):
    # Create a new Document
    doc = Document()

    # Add a title to the document
    doc.add_heading(f'Отчет о погоде за {MONTH}.{YEAR}', level=1)

    # Dictionary for translation
    translation = {
        'min_temperature': 'Минимальная температура',
        'max_temperature': 'Максимальная температура',
        'average_temperature': 'Средняя температура',
        'amplitude_temperature': 'Амплитуда температуры',
        'min_pressure': 'Минимальное давление',
        'max_pressure': 'Максимальное давление',
        'average_pressure': 'Среднее давление',
    }

    # Add data from the dictionary
    for key, value in data.items():
        # If the value is another dictionary, format it as a table
        if isinstance(value, dict):
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Параметр'
            hdr_cells[1].text = 'Значение'
            for sub_key, sub_value in value.items():
                # Translate sub_key if it's in the translation dictionary
                sub_key_translated = translation.get(sub_key, sub_key)
                row_cells = table.add_row().cells
                row_cells[0].text = sub_key_translated
                row_cells[1].text = str(sub_value)
            doc.add_paragraph()  # Add a blank line after the table
        else:
            # Translate key if it's in the translation dictionary
            key_translated = translation.get(key, key)
            # Add a paragraph for non-dictionary values
            doc.add_paragraph(f"{key_translated}: {value}")

    # Save the document
    doc.save(filename)



def get_weather(full_weather):
    full_weather = full_weather.lower()
    if "снег" in full_weather:
        return "Снег"

    if "дождь" in full_weather or "морось" in full_weather:
        return "Дождь"

    if "град" in full_weather:
        return "Град"

    return "Без Осадков"

def parse_day_element(day_element, time=6):
    element_class = "d-inline-block"
    daytimes_data = day_element.find_all("div", {"class": element_class })
    daytime_element = daytimes_data[time]
    daytime_element_raws = daytime_element.find_all("div")
    day_data = {}

    #Погода
    weather = daytime_element_raws[1].find("img")["alt"]


    #Температура
    temperature = int(daytime_element_raws[2].find("span").text.strip())


    #Направление ветра
    wind_direction = daytime_element_raws[11].text.strip()
    if wind_direction not in WIND_DIRECTIONS:
        wind_direction = "ШTЛ"

    #Давление
    pressure = int(daytime_element_raws[5].text.strip())

    day_data["Погода"] = weather
    day_data["Температура"] = temperature
    day_data["Направление ветра"] = wind_direction
    day_data["Давление"] = pressure
    return day_data



def get_stat_data(month_data):
    month_length = len(month_data)

    max_pressure = 0
    min_pressure = 10000
    sum_pressure = 0

    max_temperature = -1000
    min_temperature = 1000
    sum_temperature = 0

    wind_directions_count = {}
    weather_count = {}

    for day_data in month_data:
        # Даление
        pressure = day_data["Давление"]
        if max_pressure < pressure:
            max_pressure = pressure
        if min_pressure > pressure:
            min_pressure = pressure
        sum_pressure += pressure

        # Температура
        temperature = day_data["Температура"]
        if max_temperature < temperature:
            max_temperature = temperature
        if min_temperature > temperature:
            min_temperature = temperature
        sum_temperature += temperature

        # Ветер
        wind_direction = day_data["Направление ветра"]
        if wind_direction in wind_directions_count:
            wind_directions_count[wind_direction] += 1
        else:
            wind_directions_count[wind_direction] = 1

        #Осадки
        weather = get_weather(day_data["Погода"])
        if weather in weather_count:
            weather_count[weather] += 1
        else:
            weather_count[weather] = 1

    average_pressure = round(sum_pressure / month_length, 2)

    average_temperature = round(sum_temperature / month_length, 2)
    amplitude_temperature = max_temperature + min_temperature

    stat_data = {
        "min_temperature": min_temperature,
        "max_temperature": max_temperature,
        "average_temperature": average_temperature,
        "amplitude_temperature": amplitude_temperature,

        "min_pressure": min_pressure,
        "max_pressure": max_pressure,
        "average_pressure": average_pressure,

        "wind_directions_count": wind_directions_count,

        "weather_count": weather_count,
    }
    return stat_data

def main():
    url = f"{PAGE_URL}{YEAR}/{MONTH}"
    r = requests.get(url)

    r.raise_for_status()


    html_doc = r.text
    parser = 'html.parser'
    soup = BeautifulSoup(html_doc, parser)
    element_class = "swiper-slide swiper-autoheight w-auto d-inline-block border-start"
    days_elements = soup.find_all("div", {"class": element_class })
    month_data = []
    for day_element in days_elements:
        try:
            day_data = parse_day_element(day_element)
            month_data.append(day_data)
        except Exception:
            break

    stat_data = get_stat_data(month_data)
    generate_word_doc(stat_data, "weather_report.docx")

if __name__ == "__main__":
    try:
        main()
    except Exception:
        print("Погода за этот период не найдена")
